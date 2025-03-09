import streamlit as st
from io import BytesIO
from datetime import datetime
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# 🎨 Interface utilisateur
st.title("📄 Générateur de Documents PDF & DOCX")
st.write("Remplissez les informations ci-dessous pour générer un document.")

# 📝 Formulaire de saisie
titre = st.text_input("📌 Titre du document", "Mon Document")
date = st.date_input("📅 Date du document", datetime.today())
contenu = st.text_area("📝 Contenu", "Écrivez ici le contenu du document...")

# 📌 Choix du format (PDF ou DOCX)
format_choice = st.selectbox("📄 Format du document", ["PDF", "DOCX"])

# Fonction pour générer un document PDF
def generate_pdf(titre, date, contenu):
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    
    pdf.setTitle(titre)
    pdf.setFont("Helvetica", 12)
    
    pdf.drawString(100, 750, f"Titre: {titre}")
    pdf.drawString(100, 730, f"Date: {date}")
    text = pdf.beginText(100, 700)
    
    for line in contenu.split("\n"):
        text.textLine(line)
    
    pdf.drawText(text)
    pdf.showPage()
    pdf.save()
    
    buffer.seek(0)
    return buffer

# Fonction pour générer un document DOCX
def generate_docx(titre, date, contenu):
    doc = Document()
    doc.add_heading(titre, level=1)
    doc.add_paragraph(f"Date: {date}")
    doc.add_paragraph(contenu)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# 🔘 Bouton de génération
if st.button("🚀 Générer le document"):
    if titre and contenu:
        if format_choice == "PDF":
            file_buffer = generate_pdf(titre, date, contenu)
            st.success("✅ PDF généré avec succès !")
            st.download_button("📥 Télécharger le PDF", file_buffer, file_name="document.pdf", mime="application/pdf")

        elif format_choice == "DOCX":
            file_buffer = generate_docx(titre, date, contenu)
            st.success("✅ DOCX généré avec succès !")
            st.download_button("📥 Télécharger le DOCX", file_buffer, file_name="document.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    else:
        st.warning("⚠️ Veuillez remplir tous les champs !")

st.write("---")
st.write("🚀 **Déployé avec Streamlit**")
