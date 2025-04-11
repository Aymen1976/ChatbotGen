from flask import Flask, request, jsonify, send_from_directory
from datetime import datetime
import os
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document

app = Flask(__name__, static_url_path='/static')

# Utilisons un dossier "static" pour stocker les fichiers générés
STATIC_FOLDER = "static"
os.makedirs(STATIC_FOLDER, exist_ok=True)

# Fonction pour générer le PDF avec Platypus
def generate_pdf_platypus(titre, date, contenu, output_path):
    # Créer le document avec des marges standard
    doc = SimpleDocTemplate(output_path, pagesize=A4,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    story = []

    # Ajouter le titre avec le style "Title"
    story.append(Paragraph(titre, styles['Title']))
    story.append(Spacer(1, 12))  # Espace de 12 points

    # Ajouter la date avec un style de sous-titre
    story.append(Paragraph(f"Date: {date}", styles['Heading2']))
    story.append(Spacer(1, 12))

    # Ajouter le contenu avec le style "BodyText"
    story.append(Paragraph(contenu, styles['BodyText']))

    doc.build(story)

# Fonction pour générer le DOCX (comme avant)
def generate_docx(titre, date, contenu, output_path):
    doc = Document()
    doc.add_heading(titre, level=1)
    doc.add_paragraph(f"Date : {date}")
    doc.add_paragraph(contenu)
    doc.save(output_path)

@app.route("/")
def accueil():
    return "Bienvenue sur mon API ChatbotGen 🚀. Utilise /generer pour créer des documents."

@app.route("/generer", methods=["POST"])
def generer_documents():
    try:
        data = request.get_json()
        titre = data.get("titre", "Document sans titre")
        contenu = data.get("contenu", "")
        date = data.get("date", datetime.now().strftime("%Y-%m-%d"))

        # Vérifier que le contenu n'est pas vide
        if not contenu.strip():
            return jsonify({"error": "Le champ 'contenu' est requis"}), 400

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nom_pdf = f"{titre.replace(' ', '_')}_{timestamp}.pdf"
        nom_docx = f"{titre.replace(' ', '_')}_{timestamp}.docx"
        chemin_pdf = os.path.join(STATIC_FOLDER, nom_pdf)
        chemin_docx = os.path.join(STATIC_FOLDER, nom_docx)

        # Générer le PDF avec Platypus
        generate_pdf_platypus(titre, date, contenu, chemin_pdf)
        # Générer le DOCX
        generate_docx(titre, date, contenu, chemin_docx)

        base_url = request.host_url.rstrip("/")  # Ex: "https://chatbotgen-api.onrender.com"

        return jsonify({
            "message": "✅ Fichiers générés avec succès",
            "pdf": f"{base_url}/telecharger/pdf/{nom_pdf}",
            "docx": f"{base_url}/telecharger/docx/{nom_docx}"
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/telecharger/pdf/<nom_fichier>")
def telecharger_pdf(nom_fichier):
    # Ceci force le téléchargement du fichier
    return send_from_directory(STATIC_FOLDER, nom_fichier, as_attachment=True)

@app.route("/telecharger/docx/<nom_fichier>")
def telecharger_docx(nom_fichier):
    return send_from_directory(STATIC_FOLDER, nom_fichier, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)
