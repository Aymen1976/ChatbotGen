from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from reportlab.lib.utils import simpleSplit
from docx import Document

def generate_pdf(titre, date, contenu, output_path):
    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4
    margin = 2 * cm
    line_height = 14
    y = height - margin

    c.setFont("Helvetica-Bold", 14)
    c.drawString(margin, y, titre)
    y -= 2 * line_height

    c.setFont("Helvetica", 12)
    c.drawString(margin, y, f"Date: {date}")
    y -= 2 * line_height

    # Gestion du contenu long avec découpage automatique
    lines = simpleSplit(contenu, "Helvetica", 12, width - 2 * margin)
    for line in lines:
        if y < margin + line_height:
            c.showPage()
            y = height - margin
            c.setFont("Helvetica", 12)
        c.drawString(margin, y, line)
        y -= line_height
    c.save()

def generate_docx(titre, date, contenu, output_path):
    doc = Document()
    doc.add_heading(titre, level=1)
    doc.add_paragraph(f"Date : {date}")
    doc.add_paragraph(contenu)
    doc.save(output_path)
